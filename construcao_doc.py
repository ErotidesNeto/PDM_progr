import json
from docx import Document
import os

# Definição dos caminhos dos arquivos modelo e minuta
# Caminho do arquivo modelo.docx na pasta 'modelos'
#modelo = os.path.join("modelos", "modelo.docx")
modelo = os.path.join("modelos", "modelo.docx")
# Caminho do arquivo contestacao.docx na pasta 'documentos'
minuta = os.path.join("documentos", "contestacao.docx")

# Funções para criar e manipular o arquivo docx

def substituir_trechos_multiplos_paragrafos(doc_path, str_inicial, str_final, texto_substituto, save_path):
    """
    Substitui um trecho de texto que pode se estender por múltiplos parágrafos.
    """
    # Abre o documento especificado pelo caminho 'doc_path'
    doc = Document(doc_path)
    dentro_do_trecho = False  # Flag para indicar se está dentro do trecho que deve ser substituído
    novo_paragrafo = None  # Variável para armazenar o novo parágrafo que será criado
    runs_to_remove = []  # Lista de runs (trechos de texto com formatação) a serem removidos

    # Itera sobre todos os parágrafos do documento
    for paragrafo in doc.paragraphs:
        # Se o parágrafo contém a string inicial, indica o início do trecho a ser substituído
        if str_inicial in paragrafo.text:
            dentro_do_trecho = True
            # Cria um novo parágrafo antes do atual para iniciar a substituição
            novo_paragrafo = paragrafo.insert_paragraph_before(paragrafo.text)
            runs_to_remove = []
            # Itera sobre os runs do parágrafo para substituir o trecho desejado
            for run in paragrafo.runs:
                if str_inicial in run.text:
                    # Manipulando o texto para substituir após o str_inicial
                    inicio = run.text.find(str_inicial) + len(str_inicial)
                    texto_inicial = run.text[:inicio]
                    # Adiciona o texto inicial ao novo parágrafo, mantendo a formatação
                    novo_paragrafo.add_run(texto_inicial).bold = run.bold
                    # Adiciona o texto substituto ao novo parágrafo, mantendo a formatação
                    novo_paragrafo.add_run(texto_substituto).bold = run.bold
                    # Atualiza o texto do run para o restante após a string inicial
                    run.text = run.text[inicio:]
                    break
                else:
                    # Adiciona o texto do run ao novo parágrafo, mantendo a formatação
                    novo_paragrafo.add_run(run.text).bold = run.bold
                    runs_to_remove.append(run)
            continue

        # Se estiver dentro do trecho a ser substituído
        if dentro_do_trecho:
            if str_final in paragrafo.text:
                # Encontrou o final do trecho a ser substituído
                fim = paragrafo.text.find(str_final)
                novo_paragrafo.add_run(paragrafo.text[:fim]).bold = paragrafo.runs[0].bold
                novo_paragrafo.add_run(paragrafo.text[fim:]).bold = paragrafo.runs[0].bold
                dentro_do_trecho = False  # Encerra a substituição do trecho
            else:
                # Adiciona os runs do parágrafo atual à lista de runs a serem removidos
                for run in paragrafo.runs:
                    runs_to_remove.append(run)
            continue

        # Copia o parágrafo original, inserindo-o antes do atual
        novo_paragrafo = paragrafo.insert_paragraph_before(paragrafo.text)
        for run in paragrafo.runs:
            novo_paragrafo.add_run(run.text).bold = run.bold

    # Removendo os runs marcados
    for run in runs_to_remove:
        run.clear()
        
    # Salvando o documento modificado
    doc.save(save_path)

def substituir_trechos_mesmo_paragrafo(doc_path, str_inicial, str_final, texto_substituto, save_path):
    """
    Substitui um trecho dentro do mesmo parágrafo entre dois delimitadores.
    """
    # Abre o documento especificado pelo caminho 'doc_path'
    doc = Document(doc_path)
    
    # Itera sobre todos os parágrafos do documento
    for paragrafo in doc.paragraphs:
        # Verifica se o parágrafo contém tanto a string inicial quanto a final
        if str_inicial in paragrafo.text and str_final in paragrafo.text:
            novo_texto = ""
            inicio = 0
            # Substitui cada ocorrência do trecho entre str_inicial e str_final
            while True:
                start_idx = paragrafo.text.find(str_inicial, inicio)
                if start_idx == -1:
                    break
                end_idx = paragrafo.text.find(str_final, start_idx + len(str_inicial))
                if end_idx == -1:
                    break
                novo_texto += paragrafo.text[inicio:start_idx] + str_inicial + texto_substituto + str_final
                inicio = end_idx + len(str_final)

            novo_texto += paragrafo.text[inicio:]
            paragrafo.text = novo_texto

    # Salvando o documento modificado
    doc.save(save_path)

def substituir_string(doc_path, string_alvo, texto_substituto, save_path):
    """
    Substitui todas as ocorrências de uma string específica.
    """
    # Abre o documento especificado pelo caminho 'doc_path'
    doc = Document(doc_path)
    
    # Itera sobre todos os parágrafos do documento
    for paragrafo in doc.paragraphs:
        if string_alvo in paragrafo.text:
            # Itera sobre os runs do parágrafo e substitui a string alvo pelo texto substituto
            for run in paragrafo.runs:
                if string_alvo in run.text:
                    run.text = run.text.replace(string_alvo, texto_substituto)
    
    # Salvando o documento modificado
    doc.save(save_path)

def excluir_paragrafo_com_string(doc_path, string_alvo, save_path):
    """
    Exclui parágrafos que contêm uma string específica.
    """
    # Abre o documento especificado pelo caminho 'doc_path'
    doc = Document(doc_path)
    
    # Coleta todos os parágrafos que contêm a string alvo
    paragrafos_para_remover = [paragrafo for paragrafo in doc.paragraphs if string_alvo in paragrafo.text]
    
    # Remove cada parágrafo que contém a string alvo
    for paragrafo in paragrafos_para_remover:
        p_element = paragrafo._element
        p_element.getparent().remove(p_element)

    # Salvando o documento modificado
    doc.save(save_path)

def excluir_paragrafos_entre_strings(doc_path, str_inicial, str_final, save_path):
    """
    Exclui parágrafos que estão entre duas strings delimitadoras, incluindo os delimitadores.
    """
    # Abre o documento especificado pelo caminho 'doc_path'
    doc = Document(doc_path)
    dentro_do_trecho = False  # Flag para indicar se está dentro do trecho a ser excluído
    paragrafos_para_remover = []  # Lista de parágrafos a serem removidos

    # Itera sobre todos os parágrafos do documento
    for paragrafo in doc.paragraphs:
        if str_inicial in paragrafo.text:
            dentro_do_trecho = True
            paragrafos_para_remover.append(paragrafo)
            continue

        if dentro_do_trecho:
            paragrafos_para_remover.append(paragrafo)
            if str_final in paragrafo.text:
                dentro_do_trecho = False
            continue

    # Remove cada parágrafo que foi identificado
    for paragrafo in paragrafos_para_remover:
        p_element = paragrafo._element
        p_element.getparent().remove(p_element)

    # Salvando o documento modificado
    doc.save(save_path)

def eliminar_linhas_finais_em_branco(path_modelo, path_documento_novo):
    """
    Remove parágrafos em branco no final do documento.
    """
    # Abre o documento especificado pelo caminho 'path_modelo'
    doc = Document(path_modelo)
    
    # Remove parágrafos vazios no final do documento
    while len(doc.paragraphs) > 0 and not doc.paragraphs[-1].text.strip():
        p_element = doc.paragraphs[-1]._element
        p_element.getparent().remove(p_element)

    # Salvando o documento modificado
    doc.save(path_documento_novo)

# Funções de manipulação do arquivo de acordo com os seletores    
   

def exclui_flags(nome_flag, path_modelo, path_documento_novo):
    """
    Exclui os parágrafos que contêm as flags iniciais e finais.
    """
    # Remove parágrafos que contêm as flags iniciais e finais
    excluir_paragrafo_com_string(path_modelo, f"<FLAG_INICIAL: {nome_flag}>", path_documento_novo)
    path_modelo=path_documento_novo
    excluir_paragrafo_com_string(path_modelo, f"<FLAG_FINAL: {nome_flag}>", path_documento_novo)

def exclui_paragrafos_flag(nome_flag, path_modelo, path_documento_novo):
    """
    Exclui parágrafos que estão entre flags delimitadoras, preservando espaços.
    """
    # Utiliza a função excluir_paragrafos_entre_strings para remover parágrafos entre duas flags
    excluir_paragrafos_entre_strings(path_modelo, f"<FLAG_INICIAL: {nome_flag}>", f"<FLAG_FINAL: {nome_flag}>", path_documento_novo)    



def criar_novo_documento(analise_processual: dict, path_modelo=modelo, path_documento_novo=minuta):
    """
    Cria um novo documento com base em um modelo e análise processual.
    A análise processual é passada como um dicionário com as chaves 'prova_vinculo', 'prazo_quinquenal_inatividade' e 'servidor_temporario_nao_estavel', que trarão valores booleanos, e 'portaria' com os dados da portaria de admissão do servidor.

    """
    if analise_processual['prova_vinculo']:
        exclui_paragrafos_flag('PROVA_VINCULO', path_modelo, path_documento_novo)
        path_modelo=path_documento_novo
    else:
        exclui_flags('PROVA_VINCULO', path_modelo, path_documento_novo)
        path_modelo=path_documento_novo

    
    if not analise_processual['prazo_quinquenal_inatividade']:
        exclui_paragrafos_flag('PRAZO_QUINQUENAL_INATIVIDADE', path_modelo, path_documento_novo)
        path_modelo=path_documento_novo
    else:
        exclui_flags('PRAZO_QUINQUENAL_INATIVIDADE', path_modelo, path_documento_novo)
        path_modelo=path_documento_novo


    if not analise_processual['servidor_temporario_nao_estavel']:
        exclui_paragrafos_flag('SERVIDOR_TEMPORARIO_NAO_ESTAVEL', path_modelo, path_documento_novo)
        path_modelo=path_documento_novo
    else:
        exclui_flags('SERVIDOR_TEMPORARIO_NAO_ESTAVEL', path_modelo, path_documento_novo)
        path_modelo=path_documento_novo

    if  analise_processual['portaria']:
        substituir_string(path_modelo, '[PORTARIA]', analise_processual['portaria'], path_documento_novo)
    # TODO: Implementar a lógica para manipular o documento de acordo com a análise processual
    eliminar_linhas_finais_em_branco(path_modelo, path_documento_novo)
    
    return path_documento_novo
# Melhorias gerais:
# 1. Adicionados comentários para explicar a funcionalidade de cada função.
# 2. Removidas as funções duplicadas e renomeadas para tornar o código mais limpo e compreensível.
# 3. Otimizações na remoção e substituição de parágrafos, mantendo o código mais eficiente.
# 4. Funções reutilizáveis para evitar duplicidade de código e aumentar a manutenção.
#FLAGS: PROVA_VÍNCULO (a parte comprovou o vínculo=True), PRAZO_QUINQUENAL_INATIVIDADE (já se passaram 5 anos desde a passagem para inatividade=True), SERVIDOR_TEMPORARIO_NAO_ESTAVEL (o servidor não é estável ou é estatutário=True)
#CAMPOS: [PORTARIA] --> Indicar portaria de admissão do servidor